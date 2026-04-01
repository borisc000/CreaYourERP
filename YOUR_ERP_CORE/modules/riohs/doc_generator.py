"""
Generador de DOCX para el Reglamento Interno de Orden, Higiene y Seguridad
Basado en el formato oficial ACHS 2026 (DS N°44 / Ley 16.744 / Código del Trabajo)
"""
from __future__ import annotations
import json
import os
import tempfile
from datetime import datetime
from typing import Any, List, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.time_utils import utc_strftime


# ─────────────────────────────────────────────────────────────
# Helpers de formato
# ─────────────────────────────────────────────────────────────

def _set_heading(doc: Document, text: str, level: int = 1, color: tuple = (0, 70, 127)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p


def _add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False,
                   align: str = "left", space_before: int = 0, space_after: int = 6) -> Any:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_article(doc: Document, num: int, texto: str):
    """Agrega un artículo con número en negrita y texto justificado"""
    p = doc.add_paragraph()
    run_num = p.add_run(f"ARTÍCULO {num}°.- ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_content = p.add_run(texto)
    run_content.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_bullet(doc: Document, items: List[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)


def _json_list(val: Optional[str]) -> List[str]:
    if not val:
        return []
    try:
        result = json.loads(val)
        if isinstance(result, list):
            return result
        return [str(result)]
    except Exception:
        return [v.strip() for v in val.split(",") if v.strip()]


def _bool(val) -> bool:
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.lower() in ("true", "1", "si", "sí", "yes")
    return bool(val)


# ─────────────────────────────────────────────────────────────
# GENERADOR PRINCIPAL
# ─────────────────────────────────────────────────────────────

def generar_reglamento(cfg) -> str:
    """
    Genera el DOCX del reglamento y devuelve la ruta del archivo generado.
    """
    doc = Document()

    # Configurar márgenes (2.5cm todos)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    empresa   = cfg.empresa_nombre or "LA EMPRESA"
    rut       = cfg.empresa_rut or ""
    giro      = cfg.empresa_giro or ""
    direccion = cfg.empresa_direccion or ""
    ciudad    = cfg.empresa_ciudad or ""
    region    = cfg.empresa_region or ""
    organismo = cfg.organismo_admin or "ACHS"
    num_trab  = int(cfg.num_trabajadores or 0)
    tipo      = cfg.tipo_reglamento or ("RIOHS" if num_trab >= 10 else "RIHS")
    fecha_vig = cfg.fecha_vigencia or utc_strftime("%d/%m/%Y")
    resp_sst  = cfg.responsable_sst_nombre or "El/La Responsable de SST"
    resp_cargo = cfg.responsable_sst_cargo or "Prevencionista"
    resp_email = cfg.responsable_sst_email or ""

    jornada_hrs     = int(cfg.jornada_horas_semanales or 44)
    jornada_dias    = cfg.jornada_dias or "Lunes a Viernes"
    hora_inicio     = cfg.jornada_hora_inicio or "08:00"
    hora_fin        = cfg.jornada_hora_fin or "17:00"
    tiene_turnos    = _bool(cfg.tiene_turnos)
    desc_turnos     = cfg.descripcion_turnos or ""
    tiene_tele      = _bool(cfg.tiene_teletrabajo)

    rem_periodo     = cfg.remuneracion_periodo or "mensual"
    rem_dia         = int(cfg.remuneracion_dia or 30)
    rem_metodo      = cfg.remuneracion_metodo or "depósito bancario"
    escalas         = cfg.escalas_cargos or ""

    riesgos_fis     = _json_list(cfg.riesgos_fisicos)
    riesgos_qui     = _json_list(cfg.riesgos_quimicos)
    riesgos_bio     = _json_list(cfg.riesgos_biologicos)
    riesgos_erg     = _json_list(cfg.riesgos_ergonomicos)
    riesgos_psi     = _json_list(cfg.riesgos_psicosociales)
    epp_lista       = _json_list(cfg.epp_requeridos)
    vacunas_lista   = _json_list(cfg.vacunas_requeridas)

    alt     = _bool(cfg.trabaja_alturas)
    elec    = _bool(cfg.trabaja_electricidad)
    quim    = _bool(cfg.trabaja_quimicos)
    maq     = _bool(cfg.trabaja_maquinaria)
    esp_con = _bool(cfg.trabaja_espacios_confinados)
    pub     = _bool(cfg.trabaja_con_publico)

    multa_min = int(cfg.multa_min_pct or 1)
    multa_max = int(cfg.multa_max_pct or 25)
    email_rec = cfg.reclamos_email or "reclamos@empresa.cl"
    plazo_rec = int(cfg.reclamos_plazo or 10)

    tiene_comite = _bool(cfg.tiene_comite_paritario)
    tiene_delegado = _bool(cfg.tiene_delegado_sst)
    tiene_dpto = _bool(cfg.tiene_dpto_prevencion)

    # ────────────────────────────
    # PORTADA
    # ────────────────────────────
    doc.add_paragraph("")
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.add_run(f"{empresa}").bold = True

    doc.add_paragraph("")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_title.add_run("REGLAMENTO INTERNO DE ")
    r1.bold = True; r1.font.size = Pt(20)
    if tipo == "RIOHS":
        r2 = p_title.add_run("ORDEN, HIGIENE Y SEGURIDAD")
    else:
        r2 = p_title.add_run("HIGIENE Y SEGURIDAD")
    r2.bold = True; r2.font.size = Pt(20)

    doc.add_paragraph("")
    for line in [f"RUT: {rut}", f"Giro: {giro}", f"Dirección: {direccion}, {ciudad}", f"Región: {region}",
                 f"Organismo Administrador: {organismo}", f"N° Trabajadores: {num_trab}", f"Vigente desde: {fecha_vig}"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ────────────────────────────
    # TÍTULO 1 – ORDEN (solo si RIOHS con 10+ trabajadores)
    # ────────────────────────────
    if tipo == "RIOHS":
        _set_heading(doc, "TÍTULO 1 – REGLAMENTO INTERNO DE ORDEN", level=1)
        _add_paragraph(doc, "PREÁMBULO", bold=True, align="center")
        _add_paragraph(doc,
            f"El presente Reglamento Interno regula los requisitos y condiciones de ingreso, derechos, obligaciones, "
            f"prohibiciones y modalidades de trabajo de todas las personas trabajadoras de {empresa}, cualquiera sea "
            f"el establecimiento en que presten sus servicios. Está fundado en el C°digo del Trabajo y la Ley N° 16.744.",
            align="justify")

        # Capítulo I
        _set_heading(doc, "Capítulo I – Normas Generales", level=2)
        _add_article(doc, 1,
            f"El presente Reglamento Interno regula los requisitos y las condiciones de ingreso, derechos, obligaciones, "
            f"prohibiciones, normas, instrucciones y modalidades de trabajo de todas las personas trabajadoras de {empresa}.")
        _add_article(doc, 2,
            f"Este reglamento se entregará formalmente en una charla de inducción al personal contratado, una vez "
            f"iniciadas sus funciones, dejándose registro de ello.")
        _add_article(doc, 3,
            f"El presente reglamento es de conocimiento obligatorio y estricto cumplimiento para todas las personas "
            f"trabajadoras de {empresa}. Ningún trabajador podrá aducir desconocimiento de sus disposiciones.")
        _add_article(doc, 4,
            f"Las facultades de dirección y disciplina de {empresa} tienen como límite el respeto a las garantías "
            f"constitucionales de las personas trabajadoras, en especial su intimidad, vida privada y honra.")

        # Capítulo II - Terminología
        _set_heading(doc, "Capítulo II – Terminología y Definiciones", level=2)
        _add_article(doc, 5,
            "Para efectos del presente Reglamento se entenderá por: Accidente del Trabajo: toda lesión que una persona "
            "sufra a causa del trabajo y le produzca incapacidad o muerte (Art. 5° Ley 16.744). Enfermedad Profesional: "
            "causada directamente por el ejercicio de la profesión (Art. 7° Ley 16.744). Empleador: quien utiliza los "
            f"servicios de trabajadores en virtud de un contrato (Art. 3° Código del Trabajo). Entidad Empleadora: {empresa}. "
            "Comité Paritario: organismo mixto de empleador y trabajadores para prevención de riesgos (Art. 66 Ley 16.744). "
            "EPP: Elemento de Protección Personal. Teletrabajo: servicios prestados con medios tecnológicos desde domicilio "
            "u otro lugar distinto a la empresa.")

        # Capítulo III - Ingreso
        _set_heading(doc, "Capítulo III – Condiciones de Ingreso", level=2)
        _add_article(doc, 6,
            f"Las personas interesadas en ingresar a {empresa} deberán presentar: fotocopia de Cédula de Identidad, "
            "Currículum Vitae, Certificado de Nacimiento, Certificado de Antecedentes, Certificado de Situación "
            "Militar (hombres entre 18 y 45 años), certificados de estudios y cualquier otro documento que el cargo requiera.")
        _add_article(doc, 7,
            f"Sin perjuicio de las exigencias del cargo, constituye requisito esencial someterse a los exámenes "
            f"preocupacionales que {empresa} determine, según la idoneidad requerida por la función.")
        _add_article(doc, 8,
            f"Según el puesto de trabajo, la persona deberá cumplir con las medidas preventivas determinadas en la "
            f"Matriz de Identificación de Peligros y Evaluación de Riesgos (MIPER) de {empresa}.")

        # Capítulo IV - Contrato
        _set_heading(doc, "Capítulo IV – Del Contrato de Trabajo", level=2)
        _add_article(doc, 9,
            f"Si el postulante cumple los requisitos, {empresa} escriturará el Contrato de Trabajo dentro de los "
            f"plazos legales. Un ejemplar quedará en poder del trabajador y otro en poder de {empresa}.")
        _add_article(doc, 10,
            "El Contrato de Trabajo contendrá las estipulaciones previstas en el Art. 10° del Código del Trabajo: "
            "lugar y fecha, identificación de las partes, naturaleza de los servicios, remuneración, jornada y duración.")
        _add_article(doc, 11,
            f"Cualquier modificación en datos personales (domicilio, estado civil, cargas familiares, previsión) debe "
            f"comunicarse a {empresa} dentro de los 7 días hábiles siguientes al hecho que la motive.")

        # Capítulo V - Término
        _set_heading(doc, "Capítulo V – Terminación del Contrato de Trabajo", level=2)
        _add_article(doc, 12,
            "El contrato de trabajo podrá terminar por las causales establecidas en los artículos 159, 160 y 161 "
            "del Código del Trabajo, con los avisos y procedimientos que la ley establece.")

        # Capítulo VI - Jornada
        _set_heading(doc, "Capítulo VI – Jornadas de Trabajo", level=2)
        _add_article(doc, 13,
            f"La jornada ordinaria de trabajo en {empresa} será de {jornada_hrs} horas semanales, distribuidas "
            f"de {jornada_dias}, de {hora_inicio} a {hora_fin} horas, con el tiempo de colación que se establezca.")
        if tiene_turnos:
            _add_article(doc, 14,
                f"Las personas trabajadoras se regirán por el sistema de turnos descrito en el Anexo N°1 de este "
                f"Reglamento. Los turnos podrán ser modificados según las necesidades operativas de {empresa}.")
        if tiene_tele:
            _add_article(doc, 15,
                f"{empresa} podrá acordar con sus trabajadores modalidades de trabajo a distancia o teletrabajo, "
                "conforme a lo establecido en el Art. 152 quáter M y siguientes del Código del Trabajo.")
        _add_article(doc, 16,
            f"Se entiende por jornada extraordinaria la que excede del máximo legal o de la pactada contractualmente. "
            f"Las horas extraordinarias deberán pactarse por escrito y serán pagadas con el recargo legal correspondiente.")

        # Capítulo VII - Remuneraciones
        _set_heading(doc, "Capítulo VII – Remuneraciones", level=2)
        _add_article(doc, 17,
            f"La remuneración de cada persona trabajadora es aquella acordada individualmente en su contrato de trabajo. "
            f"El pago se realizará de forma {rem_periodo}, el día {rem_dia} de cada mes mediante {rem_metodo}.")
        _add_article(doc, 18,
            f"{empresa} dará cumplimiento al principio de igualdad de remuneraciones entre hombres y mujeres que "
            f"presten iguales funciones, de conformidad al Art. 62 bis del Código del Trabajo.")
        if escalas:
            _add_article(doc, 19,
                f"El detalle de cargos, funciones y escalas de remuneraciones se encuentra en el Anexo N°2 del "
                f"presente Reglamento.")

        # Capítulo VIII - Feriado
        _set_heading(doc, "Capítulo VIII – Feriado Anual y Permisos", level=2)
        _add_article(doc, 20,
            "Las personas trabajadoras con más de un año de servicio tendrán derecho a un feriado anual de 15 días "
            "hábiles, conforme al Art. 67 del Código del Trabajo. El feriado progresivo se regirá por lo dispuesto "
            "en el Art. 68 del mismo cuerpo legal.")
        _add_article(doc, 21,
            f"Se concederán permisos con goce de remuneración en los casos establecidos por ley: nacimiento de hijo "
            f"(5 días), fallecimiento de cónyuge o hijo (5 días), fallecimiento de padre o madre (3 días) y demás "
            f"permisos legales vigentes.")

        # Capítulo IX - Maternidad
        _set_heading(doc, "Capítulo IX – Protección a la Maternidad y Parentalidad", level=2)
        _add_article(doc, 22,
            "Las trabajadoras embarazadas y en período de lactancia gozarán de todos los derechos establecidos en el "
            "Título II del Libro II del Código del Trabajo, incluyendo fuero maternal, descanso pre y postnatal, "
            "permiso posnatal parental y sala cuna.")

        # Capítulo X - Igualdad
        _set_heading(doc, "Capítulo X – Igualdad de Oportunidades", level=2)
        _add_article(doc, 23,
            f"{empresa} garantiza igualdad de oportunidades para personas con discapacidad en los procesos de "
            f"selección y durante la relación laboral, conforme a la Ley N° 20.422.")

        # Capítulo XI - Reclamos
        _set_heading(doc, "Capítulo XI – Canal de Denuncias, Consultas y Reclamos", level=2)
        _add_article(doc, 24,
            f"Las personas trabajadoras podrán dirigir sus reclamos, consultas o denuncias a través del canal "
            f"oficial de {empresa}: {email_rec}. {empresa} responderá dentro de {plazo_rec} días hábiles.")
        _add_article(doc, 25,
            f"{empresa} mantendrá un procedimiento de investigación y sanción del acoso laboral y sexual, conforme "
            f"al Protocolo establecido por el DS N°44 y el Código del Trabajo.")

        # Capítulo XII - Obligaciones Empresa
        _set_heading(doc, "Capítulo XII – Obligaciones de la Entidad Empleadora", level=2)
        _add_article(doc, 26,
            f"{empresa} se obliga a: (a) cumplir las normas de higiene y seguridad; (b) proporcionar EPP "
            f"gratuitamente; (c) informar los riesgos del puesto de trabajo (Derecho a Saber); (d) mantener "
            f"condiciones ambientales adecuadas; (e) pagar cotizaciones previsionales oportunamente.")

        # Capítulo XIII - Obligaciones Trabajadores
        _set_heading(doc, "Capítulo XIII – Obligaciones de las Personas Trabajadoras", level=2)
        _add_article(doc, 27,
            f"Toda persona trabajadora de {empresa} está obligada a: (a) cumplir este reglamento; (b) asistir "
            f"puntualmente; (c) cuidar bienes e instalaciones; (d) usar correctamente los EPP entregados; "
            f"(e) reportar condiciones inseguras a su jefatura directa; (f) participar en capacitaciones de SST.")

        # Capítulo XIV - Prohibiciones
        _set_heading(doc, "Capítulo XIV – Prohibiciones", level=2)
        _add_article(doc, 28,
            f"Se prohíbe a las personas trabajadoras de {empresa}: (a) presentarse al trabajo en estado de "
            f"intemperancia o bajo efectos de estupefacientes; (b) introducir bebidas alcohólicas o drogas; "
            f"(c) hacer uso de teléfonos celulares en zonas de riesgo; (d) operar equipos sin autorización; "
            f"(e) efectuar cambios de turno sin autorización; (f) abandonar el trabajo sin justificación.")

        # Capítulo XV - Tabaco
        _set_heading(doc, "Capítulo XV – Regulación del Tabaco", level=2)
        _add_article(doc, 29,
            f"Queda prohibido fumar en todas las dependencias cerradas de {empresa}, conforme a la Ley N° 20.660. "
            f"Se habilitará un área exterior delimitada para fumadores donde la normativa lo permita.")

        # Capítulo XVI - Sanciones
        _set_heading(doc, "Capítulo XVI – Sanciones y Procedimiento de Reclamación", level=2)
        _add_article(doc, 30,
            f"Las infracciones a este Reglamento podrán ser sancionadas con: (a) amonestación verbal; "
            f"(b) amonestación escrita; (c) multa de {multa_min}% a {multa_max}% de la remuneración diaria, "
            f"conforme al Art. 67 Ley 16.744 y Art. 154 N°10 del Código del Trabajo.")
        _add_article(doc, 31,
            f"La persona trabajadora afectada por una sanción tendrá derecho a reclamar ante la Inspección del "
            f"Trabajo o la autoridad de salud competente, dentro de los plazos legales.")

        # Capítulo XVII - Seguridad General
        _set_heading(doc, "Capítulo XVII – Normas de Seguridad General", level=2)
        _add_article(doc, 32,
            f"Toda persona trabajadora de {empresa} tiene el deber y el derecho a trabajar en condiciones de "
            f"seguridad y salud. Ante un riesgo grave e inminente, tiene derecho a interrumpir sus actividades "
            f"y abandonar el lugar de trabajo, informando de inmediato a su jefatura.")

        # Capítulo XVIII - Sindicatos
        _set_heading(doc, "Capítulo XVIII – Organizaciones Sindicales", level=2)
        _add_article(doc, 33,
            f"{empresa} reconoce el derecho de las personas trabajadoras a constituir y afiliarse a organizaciones "
            f"sindicales, conforme al Título I del Libro III del Código del Trabajo.")

        # Capítulo XIX - No Discriminación
        _set_heading(doc, "Capítulo XIX – No Discriminación y Respeto", level=2)
        _add_article(doc, 34,
            f"{empresa} prohíbe cualquier acto de discriminación arbitraria, acoso laboral, acoso sexual o violencia "
            f"en el trabajo, comprometiéndose a investigar y sancionar estas conductas conforme al DS N°44 y la "
            f"legislación vigente.")

        doc.add_page_break()

    # ────────────────────────────
    # TÍTULO 2 – HIGIENE Y SEGURIDAD (obligatorio para TODOS)
    # ────────────────────────────
    art_base = 35 if tipo == "RIOHS" else 1
    _set_heading(doc, "TÍTULO 2 – REGLAMENTO INTERNO DE HIGIENE Y SEGURIDAD", level=1)
    _add_paragraph(doc,
        f"El presente Título es de cumplimiento obligatorio para todas las personas trabajadoras de {empresa}, "
        f"conforme al Art. 67° de la Ley N°16.744 y el Art. 56° del DS N°44.", align="justify")

    # Cap I - Disposiciones generales
    _set_heading(doc, "Capítulo I – Disposiciones Generales", level=2)
    _add_article(doc, art_base,
        f"{empresa} está afiliada al organismo administrador del seguro de la Ley N°16.744: {organismo}. "
        f"El presente reglamento es de cumplimiento obligatorio y se entregará gratuitamente a cada trabajador.")
    _add_article(doc, art_base+1,
        f"El responsable de Seguridad y Salud en el Trabajo es {resp_sst}, {resp_cargo}"
        + (f" ({resp_email})" if resp_email else "") + ".")
    if tiene_comite:
        _add_article(doc, art_base+2,
            f"{empresa} cuenta con Comité Paritario de Higiene y Seguridad (CPHS) conforme al Art. 66° de la Ley "
            f"16.744, el que participará en la elaboración y revisión de este reglamento.")
    if tiene_dpto:
        _add_article(doc, art_base+3,
            f"{empresa} cuenta con Departamento de Prevención de Riesgos a cargo de planificar, organizar y "
            f"supervisar la gestión preventiva conforme al DS N°44.")

    # Cap II - Obligaciones
    _set_heading(doc, "Capítulo II – Obligaciones en Materia de Seguridad", level=2)
    _add_article(doc, art_base+4,
        f"Toda persona trabajadora de {empresa} deberá: (a) cumplir las normas de higiene y seguridad; "
        f"(b) usar correctamente los EPP asignados; (c) informar inmediatamente todo accidente, incidente o "
        f"condición insegura; (d) participar en charlas y capacitaciones de SST; (e) no operar equipos sin "
        f"autorización ni entrenamiento.")

    # EPP
    if epp_lista:
        _add_article(doc, art_base+5,
            f"Los EPP de uso obligatorio en {empresa} son, según el puesto de trabajo:")
        _add_bullet(doc, epp_lista)
    else:
        _add_article(doc, art_base+5,
            f"Los EPP serán determinados según la Matriz de Identificación de Peligros y Evaluación de Riesgos "
            f"(MIPER) de {empresa} y serán entregados gratuitamente a cada trabajador.")

    # Vacunas
    if vacunas_lista:
        _add_article(doc, art_base+6,
            f"Según los riesgos biológicos identificados en la MIPER, los trabajadores expuestos deberán "
            f"contar con las siguientes vacunas:")
        _add_bullet(doc, vacunas_lista)

    # Riesgos identificados
    n = art_base+7
    _set_heading(doc, "Capítulo III – Identificación de Riesgos y Medidas Preventivas", level=2)

    all_riesgos = [
        ("Riesgos Físicos", riesgos_fis),
        ("Riesgos Químicos", riesgos_qui),
        ("Riesgos Biológicos", riesgos_bio),
        ("Riesgos Ergonómicos", riesgos_erg),
        ("Riesgos Psicosociales", riesgos_psi),
    ]
    for nombre_riesgo, lista_riesgo in all_riesgos:
        if lista_riesgo:
            _add_article(doc, n, f"{nombre_riesgo} identificados en {empresa}:")
            _add_bullet(doc, lista_riesgo)
            n += 1

    # Actividades especiales
    if any([alt, elec, quim, maq, esp_con]):
        _add_article(doc, n,
            f"En razón de las actividades específicas de {empresa}, aplican las siguientes medidas especiales:")
        especiales = []
        if alt:  especiales.append("Trabajos en altura física: uso obligatorio de arnés y línea de vida certificados")
        if elec: especiales.append("Trabajos eléctricos: solo personal autorizado, uso de EPP dieléctrico y bloqueo LOTO")
        if quim: especiales.append("Manejo de sustancias químicas: uso de Hojas de Datos de Seguridad (HDS/SDS)")
        if maq:  especiales.append("Operación de maquinaria pesada: solo operadores certificados con licencia vigente")
        if esp_con: especiales.append("Espacios confinados: protocolo de entrada con permiso de trabajo y vigía externo")
        _add_bullet(doc, especiales)
        n += 1

    if pub:
        _add_article(doc, n,
            f"Por la naturaleza del servicio, las personas trabajadoras que atienden público están expuestas a "
            f"riesgo de violencia de terceros. {empresa} implementará medidas de prevención y protocolo de "
            f"atención a incidentes de violencia en el trabajo.")
        n += 1

    # Cap IV - Accidentes
    _set_heading(doc, "Capítulo IV – Accidentes del Trabajo y Enfermedades Profesionales", level=2)
    _add_article(doc, n,
        f"Todo accidente del trabajo ocurrido en {empresa} deberá ser denunciado de inmediato a la jefatura directa "
        f"y reportado dentro de 24 horas al {organismo} mediante la Denuncia Individual de Accidente del Trabajo (DIAT). "
        f"También deberán denunciarse los accidentes de trayecto.")
    n += 1
    _add_article(doc, n,
        f"Toda enfermedad de posible origen profesional deberá ser reportada al {organismo} mediante la Denuncia "
        f"Individual de Enfermedad Profesional (DIEP), conforme al Art. 7° de la Ley 16.744.")
    n += 1
    _add_article(doc, n,
        f"{empresa} realizará la investigación de accidentes de trabajo graves o fatales conforme al DS N°44, "
        f"adoptando las medidas correctivas necesarias para evitar su repetición.")
    n += 1

    # Cap V - Prohibiciones seguridad
    _set_heading(doc, "Capítulo V – Prohibiciones en Materia de Seguridad", level=2)
    _add_article(doc, n,
        f"Queda estrictamente prohibido en {empresa}: (a) operar equipos o maquinaria sin autorización; "
        f"(b) retirar o inutilizar dispositivos de seguridad; (c) ingresar a zonas de riesgo sin autorización; "
        f"(d) consumir alcohol o drogas antes o durante la jornada; (e) desobedecer señaléticas de seguridad.")
    n += 1

    # Cap VI - Sanciones Seguridad
    _set_heading(doc, "Capítulo VI – Sanciones en Materia de Seguridad", level=2)
    _add_article(doc, n,
        f"El incumplimiento de las normas de seguridad será sancionado según el procedimiento establecido en el "
        f"Capítulo XVI del Título 1 (si aplica) o mediante multa de {multa_min}% a {multa_max}% de la "
        f"remuneración diaria, conforme al Art. 67° de la Ley 16.744.")
    n += 1
    _add_article(doc, n,
        f"La persona trabajadora sancionada podrá reclamar ante la Inspección del Trabajo o la Autoridad Sanitaria "
        f"correspondiente dentro de los plazos legales.")
    n += 1

    # ────────────────────────────
    # ANEXOS
    # ────────────────────────────
    doc.add_page_break()
    _set_heading(doc, "ANEXOS", level=1)

    if tiene_turnos and desc_turnos:
        _set_heading(doc, "ANEXO N°1 – Distribución de Jornadas y Turnos", level=2)
        _add_paragraph(doc, desc_turnos, align="justify")
    else:
        _set_heading(doc, "ANEXO N°1 – Distribución de Jornadas y Turnos", level=2)
        _add_paragraph(doc,
            f"Jornada ordinaria: {jornada_dias}, de {hora_inicio} a {hora_fin} hrs. "
            f"Total semanal: {jornada_hrs} horas.", align="justify")

    _set_heading(doc, "ANEXO N°2 – Escala de Remuneraciones y Cargos", level=2)
    if escalas:
        _add_paragraph(doc, escalas, align="justify")
    else:
        _add_paragraph(doc,
            f"Las remuneraciones se establecen individualmente en cada contrato de trabajo. "
            f"El período de pago es {rem_periodo}, el día {rem_dia} de cada mes.", align="justify")

    # ────────────────────────────
    # FIRMA
    # ────────────────────────────
    doc.add_page_break()
    _add_paragraph(doc, "FIRMA Y DISTRIBUCIÓN", bold=True, align="center", space_before=20)
    _add_paragraph(doc,
        f"El presente Reglamento Interno de {tipo} fue aprobado y entrará en vigencia el {fecha_vig}. "
        f"Se entrega copia gratuita a cada persona trabajadora conforme al Art. 56° DS N°44.", align="justify")
    doc.add_paragraph("")
    _add_paragraph(doc, f"_____________________________", align="center")
    _add_paragraph(doc, f"{empresa}", bold=True, align="center")
    _add_paragraph(doc, f"RUT: {rut}", align="center")
    doc.add_paragraph("")
    _add_paragraph(doc, f"Elaborado por: {resp_sst} – {resp_cargo}", align="center")
    _add_paragraph(doc, f"Fecha: {fecha_vig}", align="center")

    # ────────────────────────────
    # GUARDAR
    # ────────────────────────────
    output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "generated_docs")
    os.makedirs(output_dir, exist_ok=True)
    safe_name = empresa.replace(" ", "_").replace("/", "-")[:40]
    filename = f"RIOHS_{safe_name}_{utc_strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.abspath(os.path.join(output_dir, filename))
    doc.save(filepath)
    return filepath
